from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
import os

class DocxResultFormatter:
    def __init__(self, file_name):
        self.document = Document()
        self.file_name = file_name
        self._setup_chp_styles()  # New method call

    def _setup_chp_styles(self):
        # Add lighter styles for CHP venues
        style = self.document.styles.add_style('ChpHeading', 1)
        style.font.size = Pt(12)  # Smaller than regular headings
        style.font.color.rgb = RGBColor(100, 100, 100)  # Lighter gray

        style = self.document.styles.add_style('ChpVenue', 1)
        style.font.size = Pt(10)  # Smaller than regular venue text
        style.font.color.rgb = RGBColor(80, 80, 80)

    def add_heading(self, text, level=1):
        heading = self.document.add_heading(text, level=level)
        heading.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in heading.runs:
            run.font.rtl = True

    def add_paragraph(self, text, font_size=11, color=RGBColor(0, 0, 0)):
        paragraph = self.document.add_paragraph(text)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.rtl = True
            run.font.size = font_size
            run.font.color.rgb = color

    def add_hyperlink(self, paragraph, url, text):
        part = paragraph.part
        r_id = part.relate_to(url, 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink', is_external=True)
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        paragraph._p.append(hyperlink)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.rtl = True

    def add_venue(self, venue, average_price_map):
        self.add_heading(f"חנות - {venue.name} - מחיר כולל: ₪{venue.total_normalized_price(average_price_map)} ", level=2)
        paragraph = self.document.add_paragraph()
        self.add_hyperlink(paragraph, venue.url, venue.url)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.rtl = True
        self.add_heading("פריטים:", level=3)
        for item in venue.items:
            self.add_bullet_point(f"{item.name}: ₪{item.price}", url=item.url)
        if venue.missing_items:
            self.add_heading("פריטים חסרים:", level=3)
            for missing_item in venue.missing_items:
                self.add_bullet_point(f"{missing_item.name}. (מחיר ממוצע - ₪{average_price_map[missing_item.searched_name]})", font_size=10, color=RGBColor(128, 128, 128))

    def add_secondary_venue(self, venue, average_price_map, rank):
        self.add_heading(f"חנות {rank} - {venue.name} - מחיר כולל: ₪{venue.total_normalized_price(average_price_map)} ", level=3)
        paragraph = self.document.add_paragraph()
        self.add_hyperlink(paragraph, venue.url, venue.url)
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in paragraph.runs:
            run.font.rtl = True
        self.add_heading("פריטים:", level=4)
        for item in venue.items:
            self.add_bullet_point(f"{item.name}: ₪{item.price}", font_size=10, color=RGBColor(128, 128, 128), url=item.url)
        if venue.missing_items:
            self.add_heading("פריטים חסרים:", level=4)
            for missing_item in venue.missing_items:
                self.add_bullet_point(f"{missing_item.name}. (מחיר ממוצע - ₪{average_price_map[missing_item.searched_name]})", font_size=10, color=RGBColor(128, 128, 128))

    def add_bullet_point(self, text, font_size=12, color=RGBColor(0, 0, 0), url=None):
        paragraph = self.document.add_paragraph(style='ListBullet')
        paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = paragraph.add_run()
        if url:
            self.add_hyperlink(paragraph, url, text)
        else:
            run.text = text
        run.font.size = Pt(font_size)
        run.font.color.rgb = color
        run.font.rtl = True
        run.font.name = 'Arial'

    def add_most_expensive_venue(self, venue, average_price_map):
        self.add_heading("הסל היקר ביותר", level=1)
        self.add_venue(venue, average_price_map)

    def add_statistics(self, venues, average_price_map, chp_venues=None):
        num_venues = len(venues)
        total_prices = [venue.total_normalized_price(average_price_map) for venue in venues]
        avg_price = round(sum(total_prices) / num_venues, 2) if num_venues > 0 else 0
        min_price = min(total_prices) if total_prices else 0
        max_price = max(total_prices) if total_prices else 0

        # Calculate cheapest CHP venue price if available
        chp_min_price = None
        if chp_venues:
            chp_prices = [venue.total_price() for venue in chp_venues]
            if chp_prices:
                chp_min_price = min(chp_prices)

        self.add_heading(f"מספר חנויות שמצאו את כל הפריטים החיוניים: {num_venues}", level=1)
        stats_text = f"מחיר מינימלי: ₪{min_price}   |   מחיר מקסימלי: ₪{max_price}"
        if chp_min_price is not None:
            stats_text += f"   |   מחיר מינימלי באונליין: ₪{chp_min_price:.2f}"
        self.add_paragraph(stats_text, font_size=9)

    def add_chp_venues(self, chp_venues):
        if not chp_venues:
            return

        self.document.add_paragraph().add_run('\n')  # Add spacing
        self.add_heading('חנויות אונליין:', level=2)  # Using proper heading

        # Sort venues by total price
        sorted_venues = sorted(chp_venues, key=lambda v: v.total_price())
        
        for venue in sorted_venues:
            # Add venue name and website with venue style
            venue_para = self.document.add_paragraph(style='ChpVenue')
            venue_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            venue_para.add_run(f'{venue.name} - ').bold = True
            self.add_hyperlink(venue_para, venue.website_url, venue.website_url)

            # Add total price
            price_para = self.document.add_paragraph(style='ChpVenue')
            price_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            price_para.add_run(f'סה"כ: {venue.total_price():.2f} ₪').bold = True

            # Add items as bullet points with hyperlinks
            for item_name, (price, url) in venue.items.items():
                para = self.document.add_paragraph(style='ChpVenue')
                para.style.font.size = Pt(10)
                para.style.font.color.rgb = RGBColor(80, 80, 80)
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                
                # Add bullet point
                para.style = self.document.styles['ListBullet']
                text = f'{item_name}: {price:.2f} ₪'
                
                if url:
                    self.add_hyperlink(para, url, text)
                else:
                    para.add_run(text)

            self.document.add_paragraph().add_run('\n')  # Add spacing between venues

    def save(self):
        base_name, extension = os.path.splitext(self.file_name)
        attempt = 1
        while True:
            if os.path.exists(self.file_name):
                attempt += 1
                self.file_name = f"{base_name}({attempt}){extension}"
                continue
            try:
                self.document.save(self.file_name)
                break
            except PermissionError:
                attempt += 1
                self.file_name = f"{base_name}({attempt}){extension}"